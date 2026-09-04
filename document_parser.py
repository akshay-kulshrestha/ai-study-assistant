"""Text extraction for uploaded study material (PDF, DOCX, TXT)."""

import os

import pypdf
from docx import Document

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}
MAX_FILE_SIZE_BYTES = 15 * 1024 * 1024  # 15 MB


class ParseError(ValueError):
    """Raised when a file can't be read or produces no usable text."""


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file_path, file_type):
    if file_type == "pdf":
        text = _extract_pdf(file_path)
    elif file_type == "docx":
        text = _extract_docx(file_path)
    elif file_type == "txt":
        text = _extract_txt(file_path)
    else:
        raise ParseError(f"Unsupported file type: {file_type}")

    text = text.strip()
    if not text:
        raise ParseError(
            "No readable text was found in this file. Scanned/image-only PDFs aren't supported."
        )
    return text


def _extract_pdf(file_path):
    try:
        reader = pypdf.PdfReader(file_path)
    except Exception as exc:
        raise ParseError(f"Couldn't open this PDF: {exc}") from exc

    if reader.is_encrypted:
        # decrypt() doesn't raise on a wrong/empty password -- it returns a
        # status code (0/NOT_DECRYPTED on failure, 1 or 2 on success), so
        # the try/except here was structurally unable to catch the exact
        # case it exists for: a genuinely password-protected PDF. Any PDF
        # protected by a real (non-empty) password would silently fall
        # through, then fail later with the unrelated and misleading "no
        # readable text" / "scanned PDF" error instead of this one.
        try:
            result = reader.decrypt("")
        except Exception:
            result = 0
        if not result:
            raise ParseError("This PDF is password-protected and can't be read.")

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _extract_docx(file_path):
    try:
        doc = Document(file_path)
    except Exception as exc:
        raise ParseError(f"Couldn't open this DOCX file: {exc}") from exc

    parts = [p.text for p in doc.paragraphs]
    # Paragraphs alone skip table content entirely -- study notes often put
    # real content (vocab, comparisons, data) in tables, and it was being
    # silently dropped with no error or indication anything was missing.
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _extract_txt(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()
    # latin-1 decodes literally any byte sequence without ever raising
    # UnicodeDecodeError, so the loop below always "succeeds" on that
    # fallback even for genuinely binary files (e.g. someone renaming a
    # .png or .exe to .txt) -- the final ParseError was unreachable dead
    # code. A null byte is a cheap, standard signal of binary content that
    # real text essentially never contains, so check for that first rather
    # than trusting "it decoded" as proof this is actually text.
    if b"\x00" in raw:
        raise ParseError("This doesn't look like a text file.")
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ParseError("Couldn't decode this text file.")


def file_type_from_name(filename):
    return filename.rsplit(".", 1)[1].lower()
